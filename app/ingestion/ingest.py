"""Document ingestion and text extraction."""
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from loguru import logger

from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.core.settings import settings
from app.storage.service import save_image, save_image_file


def extract_text_from_image(image_data: bytes) -> str:
    """
    Extract text from image using OCR (Tesseract).
    
    Args:
        image_data: Image data as bytes
        
    Returns:
        Extracted text from image
    """
    try:
        import pytesseract
        from PIL import Image
        import io
        
        # Convert bytes to PIL Image
        image = Image.open(io.BytesIO(image_data))
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using Tesseract
        text = pytesseract.image_to_string(image, lang='eng')
        
        # Clean up text
        text = text.strip()
        
        if text:
            logger.info(f"Extracted {len(text)} characters from image via OCR")
        else:
            logger.info("No text found in image via OCR")
        
        return text
        
    except ImportError:
        logger.warning("pytesseract not installed, cannot perform OCR")
        return ""
    except Exception as e:
        logger.error(f"Failed to extract text from image: {e}")
        return ""


def process_images_with_ocr(images: List[bytes], doc_id: str) -> List[Dict[str, Any]]:
    """
    Process images with OCR and return OCR text chunks.
    
    Args:
        images: List of image data as bytes
        doc_id: Document ID
        
    Returns:
        List of OCR text chunks
    """
    ocr_chunks = []
    
    for i, image_data in enumerate(images):
        try:
            # Extract text from image
            ocr_text = extract_text_from_image(image_data)
            
            if ocr_text.strip():
                # Save image to storage
                image_url = save_image(image_data, doc_id, f"ocr_{i}", i)
                
                # Create OCR chunk
                ocr_chunk = {
                    "doc_id": doc_id,
                    "chunk_id": f"ocr_{i}",
                    "text": ocr_text,
                    "has_image": True,
                    "image_urls": [image_url] if image_url else [],
                    "is_ocr": True  # Flag to identify OCR chunks
                }
                
                ocr_chunks.append(ocr_chunk)
                logger.info(f"Created OCR chunk {i} with {len(ocr_text)} characters")
            
        except Exception as e:
            logger.error(f"Failed to process image {i} with OCR: {e}")
            continue
    
    return ocr_chunks


def detect_file_type(file_path: str) -> str:
    """
    Detect file type from file path and content.
    
    Args:
        file_path: Path to the file
        
    Returns:
        File type (pdf, docx, txt, md)
    """
    try:
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return 'pdf'
        elif suffix in ['.doc', '.docx']:
            return 'docx'
        elif suffix == '.md':
            return 'md'
        elif suffix == '.txt':
            return 'txt'
        else:
            # Try to detect by MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                if mime_type == 'application/pdf':
                    return 'pdf'
                elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                    return 'docx'
                elif mime_type == 'text/plain':
                    return 'txt'
                elif mime_type == 'text/markdown':
                    return 'md'
        
        logger.warning(f"Unknown file type for {file_path}, defaulting to txt")
        return 'txt'
        
    except Exception as e:
        logger.error(f"Failed to detect file type for {file_path}: {e}")
        return 'txt'


def extract_text_from_pdf(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text and images from PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Tuple of (text, list of image bytes)
    """
    try:
        import pdfplumber
        
        text_parts = []
        images = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"Page {page_num + 1}:\n{page_text}")
                
                # Extract images
                if hasattr(page, 'images'):
                    for img in page.images:
                        try:
                            # Extract image data
                            img_data = page.crop(img).to_image().original
                            if img_data:
                                images.append(img_data)
                        except Exception as img_e:
                            logger.warning(f"Failed to extract image from page {page_num + 1}: {img_e}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters and {len(images)} images from PDF")
        return full_text, images
        
    except ImportError:
        logger.error("pdfplumber not installed, cannot extract PDF content")
        return "", []
    except Exception as e:
        logger.error(f"Failed to extract PDF content: {e}")
        return "", []


def extract_text_from_docx(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text and images from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Tuple of (text, list of image bytes)
    """
    try:
        from docx import Document
        import zipfile
        import io
        
        doc = Document(file_path)
        text_parts = []
        images = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # Extract images from document
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                for file_info in zip_file.filelist:
                    if file_info.filename.startswith('word/media/'):
                        image_data = zip_file.read(file_info.filename)
                        images.append(image_data)
        except Exception as img_e:
            logger.warning(f"Failed to extract images from DOCX: {img_e}")
        
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters and {len(images)} images from DOCX")
        return full_text, images
        
    except ImportError:
        logger.error("python-docx not installed, cannot extract DOCX content")
        return "", []
    except Exception as e:
        logger.error(f"Failed to extract DOCX content: {e}")
        return "", []


def extract_text_from_txt(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text from TXT file.
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        Tuple of (text, empty list for images)
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        logger.info(f"Extracted {len(text)} characters from TXT file")
        return text, []
        
    except UnicodeDecodeError:
        # Try with different encodings
        encodings = ['latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                logger.info(f"Extracted {len(text)} characters from TXT file using {encoding}")
                return text, []
            except UnicodeDecodeError:
                continue
        
        logger.error(f"Failed to decode TXT file with any encoding")
        return "", []
    except Exception as e:
        logger.error(f"Failed to extract TXT content: {e}")
        return "", []


def extract_text_from_md(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text from Markdown file.
    
    Args:
        file_path: Path to MD file
        
    Returns:
        Tuple of (text, empty list for images)
    """
    try:
        import markdown
        
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to plain text
        html = markdown.markdown(md_content)
        
        # Remove HTML tags to get plain text
        import re
        text = re.sub(r'<[^>]+>', '', html)
        
        logger.info(f"Extracted {len(text)} characters from Markdown file")
        return text, []
        
    except ImportError:
        logger.warning("markdown library not installed, treating as plain text")
        return extract_text_from_txt(file_path)
    except Exception as e:
        logger.error(f"Failed to extract Markdown content: {e}")
        return "", []


def extract_text_and_images(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text and images from file based on file type.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Tuple of (text, list of image bytes)
    """
    file_type = detect_file_type(file_path)
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'md':
        return extract_text_from_md(file_path)
    else:  # txt or unknown
        return extract_text_from_txt(file_path)


def chunk_text(text: str, chunk_size: int = 1500, chunk_overlap: int = 300) -> List[str]:
    """
    Split text into chunks using LangChain RecursiveCharacterTextSplitter.
    
    Args:
        text: Input text to chunk
        chunk_size: Size of each chunk
        chunk_overlap: Overlap between chunks
        
    Returns:
        List of text chunks
    """
    try:
        if not text.strip():
            return []
        
        # Configure text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
        
    except Exception as e:
        logger.error(f"Failed to chunk text: {e}")
        return [text] if text else []


def ingest_document(file_path: str, doc_id: str) -> List[Dict[str, Any]]:
    """
    Ingest document and return list of chunks including OCR text.
    
    Args:
        file_path: Path to the document file
        doc_id: Document ID
        
    Returns:
        List of chunk dictionaries (text chunks + OCR chunks)
    """
    try:
        # Extract text and images
        text, images = extract_text_and_images(file_path)
        
        all_chunks = []
        
        # Process regular text chunks
        if text.strip():
            # Chunk the text
            text_chunks = chunk_text(text)
            
            if text_chunks:
                # Create text chunk dictionaries
                for i, chunk_text_content in enumerate(text_chunks):
                    chunk_dict = {
                        "doc_id": doc_id,
                        "chunk_id": i + 1,
                        "text": chunk_text_content,
                        "has_image": False,
                        "image_urls": [],
                        "is_ocr": False
                    }
                    all_chunks.append(chunk_dict)
        
        # Process images with OCR
        if images:
            logger.info(f"Processing {len(images)} images with OCR")
            ocr_chunks = process_images_with_ocr(images, doc_id)
            all_chunks.extend(ocr_chunks)
        
        if not all_chunks:
            logger.warning(f"No chunks created from {file_path}")
            return []
        
        logger.info(f"Ingested document {doc_id} into {len(all_chunks)} chunks ({len([c for c in all_chunks if not c.get('is_ocr', False)])} text + {len([c for c in all_chunks if c.get('is_ocr', False)])} OCR)")
        return all_chunks
        
    except Exception as e:
        logger.error(f"Failed to ingest document {file_path}: {e}")
        return []

import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from loguru import logger

from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.core.settings import settings
from app.storage.service import save_image, save_image_file


def extract_text_from_image(image_data: bytes) -> str:
    """
    Extract text from image using OCR (Tesseract).
    
    Args:
        image_data: Image data as bytes
        
    Returns:
        Extracted text from image
    """
    try:
        import pytesseract
        from PIL import Image
        import io
        
        # Convert bytes to PIL Image
        image = Image.open(io.BytesIO(image_data))
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using Tesseract
        text = pytesseract.image_to_string(image, lang='eng')
        
        # Clean up text
        text = text.strip()
        
        if text:
            logger.info(f"Extracted {len(text)} characters from image via OCR")
        else:
            logger.info("No text found in image via OCR")
        
        return text
        
    except ImportError:
        logger.warning("pytesseract not installed, cannot perform OCR")
        return ""
    except Exception as e:
        logger.error(f"Failed to extract text from image: {e}")
        return ""


def process_images_with_ocr(images: List[bytes], doc_id: str) -> List[Dict[str, Any]]:
    """
    Process images with OCR and return OCR text chunks.
    
    Args:
        images: List of image data as bytes
        doc_id: Document ID
        
    Returns:
        List of OCR text chunks
    """
    ocr_chunks = []
    
    for i, image_data in enumerate(images):
        try:
            # Extract text from image
            ocr_text = extract_text_from_image(image_data)
            
            if ocr_text.strip():
                # Save image to storage
                image_url = save_image(image_data, doc_id, f"ocr_{i}", i)
                
                # Create OCR chunk
                ocr_chunk = {
                    "doc_id": doc_id,
                    "chunk_id": f"ocr_{i}",
                    "text": ocr_text,
                    "has_image": True,
                    "image_urls": [image_url] if image_url else [],
                    "is_ocr": True  # Flag to identify OCR chunks
                }
                
                ocr_chunks.append(ocr_chunk)
                logger.info(f"Created OCR chunk {i} with {len(ocr_text)} characters")
            
        except Exception as e:
            logger.error(f"Failed to process image {i} with OCR: {e}")
            continue
    
    return ocr_chunks


def detect_file_type(file_path: str) -> str:
    """
    Detect file type from file path and content.
    
    Args:
        file_path: Path to the file
        
    Returns:
        File type (pdf, docx, txt, md)
    """
    try:
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return 'pdf'
        elif suffix in ['.doc', '.docx']:
            return 'docx'
        elif suffix == '.md':
            return 'md'
        elif suffix == '.txt':
            return 'txt'
        else:
            # Try to detect by MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                if mime_type == 'application/pdf':
                    return 'pdf'
                elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                    return 'docx'
                elif mime_type == 'text/plain':
                    return 'txt'
                elif mime_type == 'text/markdown':
                    return 'md'
        
        logger.warning(f"Unknown file type for {file_path}, defaulting to txt")
        return 'txt'
        
    except Exception as e:
        logger.error(f"Failed to detect file type for {file_path}: {e}")
        return 'txt'


def extract_text_from_pdf(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text and images from PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Tuple of (text, list of image bytes)
    """
    try:
        import pdfplumber
        
        text_parts = []
        images = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"Page {page_num + 1}:\n{page_text}")
                
                # Extract images
                if hasattr(page, 'images'):
                    for img in page.images:
                        try:
                            # Extract image data
                            img_data = page.crop(img).to_image().original
                            if img_data:
                                images.append(img_data)
                        except Exception as img_e:
                            logger.warning(f"Failed to extract image from page {page_num + 1}: {img_e}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters and {len(images)} images from PDF")
        return full_text, images
        
    except ImportError:
        logger.error("pdfplumber not installed, cannot extract PDF content")
        return "", []
    except Exception as e:
        logger.error(f"Failed to extract PDF content: {e}")
        return "", []


def extract_text_from_docx(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text and images from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Tuple of (text, list of image bytes)
    """
    try:
        from docx import Document
        import zipfile
        import io
        
        doc = Document(file_path)
        text_parts = []
        images = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # Extract images from document
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                for file_info in zip_file.filelist:
                    if file_info.filename.startswith('word/media/'):
                        image_data = zip_file.read(file_info.filename)
                        images.append(image_data)
        except Exception as img_e:
            logger.warning(f"Failed to extract images from DOCX: {img_e}")
        
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters and {len(images)} images from DOCX")
        return full_text, images
        
    except ImportError:
        logger.error("python-docx not installed, cannot extract DOCX content")
        return "", []
    except Exception as e:
        logger.error(f"Failed to extract DOCX content: {e}")
        return "", []


def extract_text_from_txt(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text from TXT file.
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        Tuple of (text, empty list for images)
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        logger.info(f"Extracted {len(text)} characters from TXT file")
        return text, []
        
    except UnicodeDecodeError:
        # Try with different encodings
        encodings = ['latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                logger.info(f"Extracted {len(text)} characters from TXT file using {encoding}")
                return text, []
            except UnicodeDecodeError:
                continue
        
        logger.error(f"Failed to decode TXT file with any encoding")
        return "", []
    except Exception as e:
        logger.error(f"Failed to extract TXT content: {e}")
        return "", []


def extract_text_from_md(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text from Markdown file.
    
    Args:
        file_path: Path to MD file
        
    Returns:
        Tuple of (text, empty list for images)
    """
    try:
        import markdown
        
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to plain text
        html = markdown.markdown(md_content)
        
        # Remove HTML tags to get plain text
        import re
        text = re.sub(r'<[^>]+>', '', html)
        
        logger.info(f"Extracted {len(text)} characters from Markdown file")
        return text, []
        
    except ImportError:
        logger.warning("markdown library not installed, treating as plain text")
        return extract_text_from_txt(file_path)
    except Exception as e:
        logger.error(f"Failed to extract Markdown content: {e}")
        return "", []


def extract_text_and_images(file_path: str) -> Tuple[str, List[bytes]]:
    """
    Extract text and images from file based on file type.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Tuple of (text, list of image bytes)
    """
    file_type = detect_file_type(file_path)
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'md':
        return extract_text_from_md(file_path)
    else:  # txt or unknown
        return extract_text_from_txt(file_path)


def chunk_text(text: str, chunk_size: int = 1500, chunk_overlap: int = 300) -> List[str]:
    """
    Split text into chunks using LangChain RecursiveCharacterTextSplitter.
    
    Args:
        text: Input text to chunk
        chunk_size: Size of each chunk
        chunk_overlap: Overlap between chunks
        
    Returns:
        List of text chunks
    """
    try:
        if not text.strip():
            return []
        
        # Configure text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
        
    except Exception as e:
        logger.error(f"Failed to chunk text: {e}")
        return [text] if text else []


def ingest_document(file_path: str, doc_id: str) -> List[Dict[str, Any]]:
    """
    Ingest document and return list of chunks including OCR text.
    
    Args:
        file_path: Path to the document file
        doc_id: Document ID
        
    Returns:
        List of chunk dictionaries (text chunks + OCR chunks)
    """
    try:
        # Extract text and images
        text, images = extract_text_and_images(file_path)
        
        all_chunks = []
        
        # Process regular text chunks
        if text.strip():
            # Chunk the text
            text_chunks = chunk_text(text)
            
            if text_chunks:
                # Create text chunk dictionaries
                for i, chunk_text_content in enumerate(text_chunks):
                    chunk_dict = {
                        "doc_id": doc_id,
                        "chunk_id": i + 1,
                        "text": chunk_text_content,
                        "has_image": False,
                        "image_urls": [],
                        "is_ocr": False
                    }
                    all_chunks.append(chunk_dict)
        
        # Process images with OCR
        if images:
            logger.info(f"Processing {len(images)} images with OCR")
            ocr_chunks = process_images_with_ocr(images, doc_id)
            all_chunks.extend(ocr_chunks)
        
        if not all_chunks:
            logger.warning(f"No chunks created from {file_path}")
            return []
        
        logger.info(f"Ingested document {doc_id} into {len(all_chunks)} chunks ({len([c for c in all_chunks if not c.get('is_ocr', False)])} text + {len([c for c in all_chunks if c.get('is_ocr', False)])} OCR)")
        return all_chunks
        
    except Exception as e:
        logger.error(f"Failed to ingest document {file_path}: {e}")
        return []





